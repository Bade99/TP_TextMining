import glob
from os.path import join, dirname, abspath
from functools import partial

import pdftotext as pdftotext
from docx import Document

get_dir = partial(join, dirname(abspath(__file__)), 'tests_data')
ds_file = partial(join, dirname(abspath(__file__)), 'dataset-nlp-utn')


def test_open_doc():
    doc = Document(get_dir('ej1.docx'))
    result = '\n'.join([x.text for x in doc.paragraphs if x.text])
    expected_result = (
        'Hola como estas\n'
        'esto es \n'
        'un archivo de \n'
        'Ejemplo\n'
        '123\n'
        '1\n'
        '412413241     afeaf'
    )
    assert result == expected_result


def test_open_pdf():
    with open(get_dir('ej1.pdf'), 'rb') as f:
        pdf = pdftotext.PDF(f)

    result = '\n'.join(pdf)
    expected_result = (
        "Hola como estas\n"
        "esto es\n"
        "un archivo de\n"
        "Ejemplo\n"
        "123\n"
        "1\n"
        "412413241     afeaf\n"
    )
    assert result == expected_result


test_open_doc()
test_open_pdf()


def doc_to_text(filename):
    return '\n'.join(
        x.text for x in
        Document(filename).paragraphs
        if x.text
    )


def pdf_to_text(filename):
    with open(filename, 'rb') as f:
        pdf = pdftotext.PDF(f)

    return '\n'.join(pdf)


def test_can_read_all_pdfs():
    for i, file in enumerate(glob.glob('./dataset-nlp-utn/*.pdf')):
        assert pdf_to_text(file)


def test_can_read_all_docs():
    for i, file in enumerate(glob.glob('./dataset-nlp-utn/*.docx')):
        assert doc_to_text(file)


test_can_read_all_docs()
test_can_read_all_pdfs()